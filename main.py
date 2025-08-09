import os
import fitz
from docx import Document as DocxDocument
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chat_models import init_chat_model
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain.memory import ConversationBufferMemory
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from typing import List
import gradio as gr

# ----------- Config -----------
if not os.environ.get("GROQ_API_KEY"):
  os.environ["GROQ_API_KEY"] = "Your_API_Key_here"  # Set your Groq API key here

HARDCODED_FOLDER_PATH = "./Insurance PDFs"  # üëà Change this path to your folder

SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.doc', '.md'}

# ----------- Loaders -----------
def load_pdf(path: str) -> List[Document]:
    doc = fitz.open(path)
    return [Document(page_content=page.get_text(), metadata={"source": path}) for page in doc]

def load_docx(path: str) -> List[Document]:
    doc = DocxDocument(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    return [Document(page_content=text, metadata={"source": path})]

def load_txt(path: str) -> List[Document]:
    with open(path, 'r', encoding='utf-8') as f:
        text = f.read()
    return [Document(page_content=text, metadata={"source": path})]

def load_file(path: str) -> List[Document]:
    ext = os.path.splitext(path)[-1].lower()
    if ext == '.pdf':
        return load_pdf(path)
    elif ext == '.docx' or ext == '.doc':
        return load_docx(path)
    elif ext == '.txt':
        return load_txt(path)
    elif ext == '.md':
        return load_txt(path)
    return []

def load_documents_from_folder(folder_path: str) -> List[Document]:
    docs = []
    for root, _, files in os.walk(folder_path):
        for file in files:
            full_path = os.path.join(root, file)
            if os.path.splitext(file)[-1].lower() in SUPPORTED_EXTENSIONS:
                try:
                    docs.extend(load_file(full_path))
                except Exception as e:
                    print(f"Error loading {file}: {e}")
    return docs

def build_vectorstore(documents: List[Document]):
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    chunks = splitter.split_documents(documents)
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")
    return FAISS.from_documents(chunks, embeddings)

# ----------- Global State -----------
db = None
memory = ConversationBufferMemory(
    memory_key="history",
    input_key="question",
    return_messages=True
)

# ----------- RAG Logic -----------
def ques_responses(question, history, system_prompt, token_limit):
    global db
    if db is None:
        return "‚ùå No documents loaded. Check backend folder path."

    retriever_docs = db.similarity_search(question, k=3)
    context = "\n".join([doc.page_content for doc in retriever_docs])

    # instruction = f"""
    #         1. If the user is greeting you then greet the user and tell your name.

    #         2. If user is asking you a question then answer the question = {question}.

    #         3. If you do not find any answer for any {question}, return "Sorry, I don't know."

    #         4. If the user asks you to summarize the context, then summarize it in a concise manner.

    #         5. If the user asks you to explain something, then explain it in a simple and clear manner. Do not repeat the greeting lines for the same session
    #         """

    instruction = f"""

            1. Answer the question: {question} based on the provided context.

            2. If you do not find any relevant information for the question in the provided documents, simply respond with "Sorry, I don't know."

            3. If the user asks you to summarize the context, then summarize it in a concise manner.

            4. If the user asks you to explain something, then explain it in a simple and clear manner.

            5. Do not include any greeting messages like "Hi, I'm Meera, chatbot assistant" in your responses. Provide direct, helpful answers without unnecessary introductions.
            """

    template = """
            {system_prompt}

            {context}

            {instruction}

            Conversation history:
            {history}

            Question: {question}

            Answer:
            """

    prompt = PromptTemplate(
        input_variables=["system_prompt", "context", "instruction", "question", "history"],
        template=template,
    )

    # Initialize the chat model
    # Note: Ensure you have the correct model name and provider
    llm = init_chat_model("llama3-8b-8192", model_provider="groq")
    chain = LLMChain(llm=llm, prompt=prompt, memory=memory)

    response = chain.predict(
        question=question,
        context=context,
        instruction=instruction,
        system_prompt=system_prompt,
    )
    return response


# ----------- Load at Startup -----------
def initialize_db():
    global db
    print(f"üìÅ Loading documents from: {HARDCODED_FOLDER_PATH}")
    if not os.path.isdir(HARDCODED_FOLDER_PATH):
        raise ValueError(f"Invalid folder: {HARDCODED_FOLDER_PATH}")

    docs = load_documents_from_folder(HARDCODED_FOLDER_PATH)
    if not docs:
        raise ValueError("No supported files found in folder.")
    db = build_vectorstore(docs)
    print(f"‚úÖ {len(docs)} chunks loaded.")


# ----------- Gradio Chat UI -----------
initialize_db()  # load everything before launching the app

gr.ChatInterface(
    fn=ques_responses,
    additional_inputs=[
        gr.Textbox("You are Meera, an assistant AI chatbot.", label="System Prompt"),
        gr.Slider(10, 300, step=10, value=100, label="Max Tokens")
    ],
    title="üí¨ Ask Meera",
    description=f"Chat with documents from: `{HARDCODED_FOLDER_PATH}`",
    theme="soft",
    # retry_btn=None,
    # undo_btn=None,
    # clear_btn="Clear",
).queue().launch()
